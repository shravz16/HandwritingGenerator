import docx
from docx.shared import Pt
from docx.shared import Inches
import os
from docx2pdf import convert
from tkinter import *
from tkinter import filedialog
import tkinter as tk 
import comtypes.client
import win32com.client
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
import string
import time
import sys
sys.setrecursionlimit(20000)



defaultLoc=os.getcwd()
root=tk.Tk()
root.title("Hand Writing generator")

root.config(background="#7DCEA0")
root.attributes('-fullscreen',True)
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()

def stop():
    root.destroy()
def select():
    gg=filedialog.askopenfilename(initialdir=defaultLoc,title="choose a docx file",filetypes=(("docx files","*docx")   ,("docx files","*docx")))
    message = tk.Label(root, text="" ,bg="#7DCEA0"  ,fg="#7DCEA0"  ,width=50  ,height=3,font=('caliber', 30, 'bold')) 
    message.place(x=(0.447916*screen_width), y=(0.64814814*screen_height))
    
    doc = docx.Document(gg)
    doc1=docx.Document()
    temp=gg
    listt=list(temp.split('/'))
    name=listt[-1]
    nname=list(name.split('.'))
    name=nname[0]
    
    


             
    start(doc,doc1,name)
    

def start(doc,doc1,name):
    wdFormatPDF=17
    
    for para in doc.paragraphs:
        paragraph = doc1.add_paragraph()
        run = paragraph.add_run(para.text)
        font = run.font
        font.name = "My Font"
        font.size = Pt(24)
    #paragraph.paragraph_format.left_indent = Inches(2.5)
    doc1.save('temp.docx')
    word = win32com.client.Dispatch('Word.Application')
    
    #doc1.close()
    doc = word.Documents.Open(defaultLoc+'/temp.docx')
    doc.SaveAs(defaultLoc+'/'+name, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    message = tk.Label(root, text="Successfully converted to "+name+".pdf" ,bg="green"  ,fg="white"  ,width=50  ,height=3,font=('caliber', 30, 'bold')) 
    message.place(x=(0.207916*screen_width), y=(0.44814814*screen_height))

    

    
    

    


root.bind("<Escape>",exit) 
message = tk.Label(root, text="Document to Handwriting" ,bg="#5DADE2"  ,fg="white"  ,width=50  ,height=3,font=('caliber', 30, 'bold')) 
message.place(x=(0.20916*screen_width),y=20)
space = tk.Button(root, text="click to start", command=select  ,fg="white"  ,bg="#2E86C1"  ,width=20  ,height=3, activebackground = "white" ,font=('times', 15, ' bold '))
space.place(x=(0.447916*screen_width), y=(0.44814814*screen_height))
close=tk.Button(root,text="X",command=stop,fg="white"  ,bg="red"  ,width=4  ,height=2, activebackground = "white" ,font=('times', 10, ' bold '))
close.place(x=(0.963541*screen_width),y=(0.00925925*screen_height))
root.mainloop()




